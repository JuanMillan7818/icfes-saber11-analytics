"""
Generador de informe Word — EduMetrics Saber11 Saber 11
Ejecutar: uv run python generar_informe.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Configuración de página ────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Cm(21.59)  # Carta / Letter
section.page_height = Cm(27.94)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)


# ── Helpers de estilo ──────────────────────────────────────────────────────────
def h1(text: str):
    p = doc.add_paragraph(text, style="Heading 1")
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    p.runs[0].font.size = Pt(18)
    p.runs[0].bold = True
    return p


def h2(text: str):
    p = doc.add_paragraph(text, style="Heading 2")
    p.runs[0].font.color.rgb = RGBColor(0x2C, 0x52, 0x82)
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True
    return p


def h3(text: str):
    p = doc.add_paragraph(text, style="Heading 3")
    p.runs[0].font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    p.runs[0].font.size = Pt(12)
    p.runs[0].bold = True
    return p


def body(text: str):
    p = doc.add_paragraph(text, style="Normal")
    p.paragraph_format.first_line_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def bullet(text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    return p


def bold_inline(p, key: str, val: str):
    """Agrega un párrafo con clave en negrita + valor normal."""
    run_k = p.add_run(key)
    run_k.bold = True
    run_k.font.size = Pt(10.5)
    run_v = p.add_run(val)
    run_v.font.size = Pt(10.5)


def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Encabezado
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A365D")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Filas de datos
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        fill = "EBF4FF" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(cell_val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tbl


def divider():
    """Línea separadora horizontal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2B6CB0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
# Logo / badge visual con texto
cover_badge = doc.add_paragraph()
cover_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_badge.add_run("📊  EduMetrics Saber11 — Saber 11")
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
cover_badge.paragraph_format.space_before = Pt(60)

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = cover_sub.add_run("Plataforma de Analítica Educativa — Microdatos 2015–2025")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

doc.add_paragraph()
divider()
doc.add_paragraph()

info_lines = [
    ("Institución: ", "INTEP — Semestre 7 · Análisis de Datos"),
    ("Proyecto: ", "Análisis de resultados ICFES Saber 11 (2015–2025)"),
    ("Tecnologías: ", "Python 3.12 · DuckDB · Streamlit · Scikit-learn · Gemini AI"),
    ("Fecha: ", datetime.date.today().strftime("%d de %B de %Y")),
]
for key, val in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bold_inline(p, key, val)

doc.add_paragraph()
divider()
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════
h1("Tabla de Contenido")
toc_items = [
    ("1.", "Introducción"),
    ("2.", "Problema y Justificación"),
    ("3.", "Pipeline ETL — Extracción, Transformación y Carga"),
    ("4.", "Vista Explorar — Análisis Exploratorio de Datos (EDA)"),
    ("  4.1", "Submódulo Análisis"),
    ("  4.2", "Submódulo Tendencias"),
    ("5.", "Herramientas (Tableros de Gestión)"),
    ("  5.1", "Coordinador Institucional"),
    ("  5.2", "Secretario de Educación"),
    ("  5.3", "Simulador Predictivo ML"),
    ("  5.4", "Priorización IPE + IA Gemini"),
    ("  5.5", "Análisis de Sensibilidad de Pesos del IPE"),
    ("6.", "Vista COVID-19 — Impacto Pandémico"),
    ("7.", "Módulo de Inteligencia Artificial (Gemini)"),
    ("8.", "Estructura Técnica del Proyecto"),
    ("9.", "Arquitectura y Dinámica de Módulos"),
    ("10.", "Conclusiones"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(title)
    r2.font.size = Pt(10.5)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Introducción")
body(
    "EduMetrics Saber11 es una plataforma modular de analítica educativa diseñada para "
    "explorar, analizar y modelar los microdatos del examen de Estado Saber 11 del ICFES, "
    "abarcando el período comprendido entre 2015 y 2025. La herramienta procesa "
    "aproximadamente 4,3 GB de datos planos, transformándolos en un conjunto de archivos "
    "Parquet comprimidos que son consultados en tiempo real mediante DuckDB."
)
body(
    "El sistema está construido bajo principios de arquitectura limpia (Clean Architecture), "
    "con una separación clara entre la capa de datos (ETL), la capa de consulta (QueryService/DuckDB) "
    "y la capa de presentación (Streamlit Dashboard). Esta separación garantiza que los módulos "
    "puedan evolucionar de forma independiente."
)
body(
    "La interfaz de usuario es un dashboard interactivo construido con Streamlit, que ofrece "
    "cinco grandes vistas: Inicio, Explorar, Herramientas, COVID-19 y Acerca de. Cada vista "
    "está orientada a un perfil específico de usuario: estudiantes, coordinadores "
    "institucionales, secretarios de educación y analistas de política pública."
)

h2("Tecnologías Principales")
tech_headers = ["Componente", "Tecnología", "Propósito"]
tech_rows = [
    ["Gestión de paquetes", "uv (Astral)", "Entorno virtual y dependencias rápidas"],
    ["Base de datos OLAP", "DuckDB", "Consultas SQL sobre Parquet in-process"],
    ["ETL y transformación", "PyArrow + Pandas", "Lectura, mapeo y escritura Parquet"],
    ["Dashboard / UI", "Streamlit", "Interfaz web interactiva sin frontend JS"],
    ["Modelos ML", "Scikit-learn", "Ridge, Random Forest, Gradient Boosting"],
    [
        "Visualizaciones",
        "Plotly Express / Graph Objects",
        "Gráficas interactivas de alta calidad",
    ],
    [
        "IA Generativa",
        "Gemini Flash (Google)",
        "Diagnósticos textuales automatizados",
    ],
]
add_table(tech_headers, tech_rows, [4, 4.5, 7])
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEMA Y JUSTIFICACIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Problema y Justificación")
body(
    "El Instituto Colombiano para la Evaluación de la Educación (ICFES) publica "
    "periódicamente los microdatos de los resultados del examen Saber 11, que constituye "
    "la principal evaluación estandarizada del nivel de educación media en Colombia. "
    "Sin embargo, estos archivos planos —distribuidos en formato .txt con separador "
    "de punto y coma— presentan múltiples barreras para su aprovechamiento efectivo:"
)
bullet(
    "Tamaño masivo: aproximadamente 4,3 GB de datos distribuidos en 22 archivos semestrales (2015-1 a 2025-2)."
)
bullet(
    "Heterogeneidad de esquemas: los nombres de columnas han cambiado a lo largo de los años, dificultando la unificación."
)
bullet(
    "Problemas de codificación: presencia de caracteres especiales, mojibake (Ã± en lugar de ñ), y variantes de encoding (UTF-8 / CP1252)."
)
bullet(
    "Ausencia de herramientas de visualización accesibles: los datos crudos no permiten a coordinadores o secretarios de educación extraer conclusiones rápidamente."
)
bullet(
    "Brecha digital no cuantificada: no existe un índice estandarizado que relacione las condiciones socioeconómicas con el rendimiento académico a nivel institucional."
)

body(
    "EduMetrics Saber11 aborda estos problemas mediante una arquitectura desacoplada que "
    "automatiza la ingesta, normalización y almacenamiento de los datos, y expone los "
    "resultados a través de visualizaciones interactivas orientadas a diferentes actores "
    "del sistema educativo. En particular, la plataforma busca:"
)
bullet(
    "Democratizar el acceso al análisis: cualquier coordinador puede consultar el rendimiento de su institución sin conocimientos de programación."
)
bullet("Cuantificar el impacto del COVID-19 sobre el rendimiento académico nacional.")
bullet(
    "Predecir puntajes futuros a partir de variables socioeconómicas mediante modelos de machine learning."
)
bullet(
    "Priorizar la intervención pública utilizando un índice compuesto (IPE) que combina deterioro académico, brecha digital y vulnerabilidad familiar."
)
bullet(
    "Generar diagnósticos textuales automáticos mediante IA generativa (Gemini) para apoyar la toma de decisiones institucionales."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ETL
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Pipeline ETL — Extracción, Transformación y Carga")

h2("3.1 Visión General del Proceso")
body(
    "El módulo ETL se encuentra en src/icfes/etl/ y está compuesto por tres archivos: "
    "config.py (mapeos de columnas), pipeline.py (ejecutor del flujo) y schemas.py "
    "(esquemas estrictos PyArrow). El proceso transforma los archivos planos originales "
    "(.txt, separados por punto y coma) en archivos .parquet comprimidos con el algoritmo "
    "Snappy, listos para ser consultados por DuckDB."
)

h2("3.2 Etapas del Pipeline")

h3("Etapa 1 — Carga de archivos fuente")
body(
    "El pipeline itera sobre todos los archivos .txt en la carpeta de datos crudos "
    "(configurable vía variable de entorno DATA_PATH_TEXT). Cada archivo corresponde "
    "a un período semestral del examen (e.g., Examen_Saber_11_20151.txt para el primer "
    "semestre de 2015). La lectura se realiza con PyArrow CSV, especificando codificación "
    "UTF-8 y delimitador de punto y coma."
)

h3("Etapa 2 — Mapeo de columnas con MAP_ATTR_ICFES")
body(
    "El diccionario MAP_ATTR_ICFES en config.py define un mapa de equivalencias entre "
    "los nombres originales de columnas de cada archivo y los nombres canónicos "
    "estandarizados del sistema. Esto permite manejar la variabilidad histórica de "
    "esquemas sin romper el pipeline. MAP_ATTR_BASE provee el mapa base que se hereda "
    "por todos los períodos, y cada período puede extenderlo (e.g., 2023 en adelante "
    "agrega estu_grupoetnia)."
)

h3("Etapa 3 — Normalización y limpieza")
body("Se aplican varias transformaciones de limpieza durante la extracción:")
bullet(
    "Normalización de área (cole_area_ubicacion): convierte variantes textuales a 'URBANO' o 'RURAL' estrictamente."
)
bullet(
    "Normalización de ubicación (departamento y municipio): aplica NFC Unicode, fix de mojibake UTF-8/CP1252, corrección de Ń→Ñ y eliminación de marcas diacríticas no relevantes."
)
bullet(
    "Corrección de nombres de instituciones: reemplaza el carácter ¿ (corrupción de Ñ en mid-word) por el carácter correcto."
)
bullet(
    "Cálculo de desempeño derivado: si la columna de desempeño no existe pero sí el puntaje, se calcula categorizando con cut en rangos [0–35, 35–50, 50–70, 70–100]."
)

h3("Etapa 4 — Enriquecimiento DIVIPOLA")
body(
    "Se carga el catálogo oficial DANE (DIVIPOLA) para enriquecer los nombres de "
    "departamentos y municipios a partir de sus códigos. Esto garantiza consistencia "
    "independientemente de la forma en que el texto original venga escrito en el archivo "
    "fuente. El catálogo se lee desde files/dane/DIVIPOLA_dane.csv."
)

h3("Etapa 5 — Inferencia de tipos y escritura Parquet")
body(
    "La función inferir_y_castear() inspecciona columna por columna: si más del 90% "
    "de los valores no nulos son numéricos, la columna se almacena como float64; "
    "en caso contrario se normaliza a string con NFKC Unicode. Finalmente, la tabla "
    "se escribe como .parquet con compresión Snappy mediante pyarrow.parquet, "
    "logrando una reducción de tamaño de aproximadamente 70-80% respecto al .txt original."
)

h2("3.3 Columnas del Esquema Maestro")
body(
    "El esquema final contiene las siguientes categorías de columnas canónicas, "
    "garantizando consistencia entre todos los períodos procesados:"
)
schema_headers = ["Prefijo", "Variables", "Descripción"]
schema_rows = [
    [
        "cole_",
        "nombre_establecimiento, depto_ubicacion, mcpio_ubicacion, naturaleza, area_ubicacion, calendario, bilingue, genero, jornada, codigo_icfes, cod_dane_establecimiento",
        "Datos del establecimiento educativo",
    ],
    [
        "estu_",
        "consecutivo, genero, fechanacimiento, depto_reside, mcpio_reside, pais_reside, nse_establecimiento, nse_individual, discapacidad, horassemanatrabaja, grupoetnia",
        "Datos del estudiante",
    ],
    [
        "fami_",
        "educacionmadre, educacionpadre, estratovivienda, tieneinternet",
        "Variables de contexto familiar",
    ],
    [
        "punt_",
        "global, matematicas, lectura_critica, c_naturales, sociales_ciudadanas, ingles",
        "Puntajes por área de conocimiento",
    ],
    [
        "desemp_",
        "matematicas, lectura_critica, c_naturales, sociales_ciudadanas, ingles",
        "Nivel de desempeño categorizado (1–4)",
    ],
    [
        "ano, periodo",
        "—",
        "Año y semestre del examen, extraídos del nombre del archivo",
    ],
]
add_table(schema_headers, schema_rows, [2.5, 8, 5])
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. EXPLORAR — EDA
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Vista Explorar — Análisis Exploratorio de Datos (EDA)")
body(
    "La pestaña Explorar del dashboard es el núcleo del análisis descriptivo. "
    "Está compuesta por dos submódulos seleccionables mediante un menú horizontal: "
    "Análisis y Tendencias. Esta vista se integra con los filtros globales del sidebar "
    "(año, departamento, género, naturaleza del colegio) que afectan las consultas "
    "DuckDB en tiempo real mediante cláusulas WHERE dinámicas."
)

h2("4.1 Submódulo Análisis")
body(
    "El submódulo Análisis (analisis.py) es la primera vista del EDA y sirve como "
    "tablero de control de KPIs generales. Ejecuta consultas DuckDB sobre el conjunto "
    "completo de Parquet, filtrado por los parámetros activos del sidebar."
)

h3("KPIs principales")
bullet("Total de estudiantes registrados en el filtro activo.")
bullet("Promedio global del puntaje ICFES (escala 0–500).")
bullet("Promedio específico de Matemáticas.")
bullet("Promedio específico de Lectura Crítica.")
bullet(
    "Departamento líder: aquel con el mayor promedio global en el período seleccionado."
)

h3("Gráficos de la sección Análisis")
bullet(
    "Tendencia anual del puntaje global: gráfico de área con relleno semitransparente, escala Y dinámica con padding adaptativo para amplificar diferencias visibles."
)
bullet(
    "Radar de competencias: gráfico polar Scatterpolar con las cinco áreas de conocimiento (Matemáticas, Lectura Crítica, Ciencias, Sociales, Inglés), permitiendo visualizar el perfil de fortalezas y debilidades."
)
bullet(
    "Mapa horizontal de departamentos: barras horizontales ordenadas por promedio descendente, con escala de color continua azul y datos hover de número de estudiantes por departamento."
)

h2("4.2 Submódulo Tendencias")
body(
    "El submódulo Tendencias (tendencias.py) provee el análisis de series temporales "
    "históricas y la proyección de puntajes futuros mediante un modelo de regresión "
    "lineal OLS."
)

h3("Serie histórica global con proyección OLS")
body(
    "Se grafica la evolución del puntaje global promedio anual desde 2015 hasta el "
    "último período disponible, junto con una proyección hacia 2026 y 2027. La "
    "proyección usa exclusivamente los datos post-COVID (2022 en adelante) para evitar "
    "que la distorsión del valle pandémico sesgue el ajuste lineal. Se ancla "
    "visualmente al último dato real para continuidad de la línea."
)
body(
    "Las métricas del modelo OLS se muestran en chips informativos: R² del ajuste, "
    "MSE, RMSE, pendiente (pts/año), número de puntos usados en el ajuste y nivel "
    "de confianza de la proyección (Alta / Moderada / Baja según umbrales de R²)."
)

h3("Series por área de conocimiento")
body(
    "Gráfico de líneas múltiples con las cinco áreas de conocimiento a lo largo "
    "de todos los años disponibles (2015–2025), con codificación de colores "
    "diferenciada para cada área. Permite comparar la evolución relativa y detectar "
    "cuáles áreas han mostrado mayor deterioro o recuperación."
)

h3("Histograma de distribución del puntaje global")
body(
    "Histograma de 60 bins del puntaje global con hasta 50,000 registros muestreados "
    "de la base de datos, mostrando la distribución de frecuencias. Cuando los datos "
    "reales no están disponibles, se usa un mock paramétrico con distribución "
    "bimodal realista (media ~255 + cola superior ~310)."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. HERRAMIENTAS
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Herramientas — Tableros de Gestión")
body(
    "La pestaña Herramientas agrupa cuatro submódulos especializados, cada uno "
    "diseñado para un perfil de usuario específico del sistema educativo colombiano. "
    "El acceso se realiza mediante un menú horizontal de cuatro opciones."
)

h2("5.1 Coordinador Institucional (coordinador.py)")
body(
    "Herramienta orientada al coordinador académico de una institución específica. "
    "Permite comparar el desempeño de su colegio contra el promedio nacional por área "
    "de conocimiento, y analizar el impacto de factores socioeconómicos."
)
bullet(
    "Selector de institución por código DANE: agrupa el historial del establecimiento aunque haya cambiado de nombre oficial, garantizando continuidad temporal."
)
bullet(
    "Gráfico divergente de brecha por área: barras horizontales rojas (por debajo del nacional) y verdes (por encima), con el cero al centro. Permite identificar de un vistazo cuáles áreas necesitan refuerzo."
)
bullet(
    "Impacto del internet en el hogar: barras agrupadas por área comparando promedios de estudiantes Con Internet, Sin Internet y con Dato desconocido."
)
bullet(
    "Impacto de horas de trabajo semanal: barras del promedio global por categoría de horas trabajadas (0h, 1-10h, 11-20h, >20h), filtrado por la institución seleccionada."
)

h2("5.2 Secretario de Educación (secretario.py)")
body(
    "Tablero diseñado para el análisis regional desde la perspectiva de equidad "
    "entre tipos de institución (Oficial vs. No Oficial) y entorno geográfico "
    "(Urbano vs. Rural)."
)
bullet("Filtros en cascada: departamento → municipio → naturaleza del establecimiento.")
bullet(
    "Gráfico doble lado a lado: tendencia temporal Oficial vs. No Oficial (izquierda) y Urbano vs. Rural (derecha) con escala Y dinámica para amplificar diferencias pequeñas."
)
bullet(
    "KPIs de rezago post-pandemia: promedio pre-COVID, durante COVID, post-COVID; caída en puntos y porcentaje de recuperación."
)

h2("5.3 Simulador Predictivo ML (simulador.py)")
body(
    "Módulo de predicción que permite proyectar el puntaje global promedio de una "
    "entidad educativa ante cambios hipotéticos en sus variables socioeconómicas. "
    "Tres modelos de machine learning corren en paralelo y sus resultados se comparan."
)

h3("Flujo de uso — Selección progresiva por pasos")
bullet("Paso 1: elegir nivel de análisis (Nacional / Municipio / Colegio).")
bullet("Paso 2 (condicional): seleccionar municipio del catálogo.")
bullet("Paso 3 (condicional): seleccionar colegio dentro del municipio.")
body(
    "Una vez seleccionada la entidad, el sistema consulta los indicadores reales "
    "actuales (dos últimos años disponibles) para precargar los sliders con valores "
    "de referencia: % con internet, % padres con educación superior y estrato promedio."
)

h3("Variables de simulación")
bullet("% Estudiantes con Internet en el hogar (0–100%).")
bullet("% Padres con educación superior — universitaria, postgrado o técnica (0–100%).")
bullet("Estrato socioeconómico promedio (1.0–6.0).")

h3("Resultados en paralelo")
body(
    "Los tres modelos predicen simultáneamente y sus resultados se muestran en "
    "tarjetas visuales. El modelo con mayor R² Test queda destacado con borde de color, "
    "glow visual y badge '🏆 Mejor modelo'. Cada tarjeta muestra el puntaje predicho, "
    "el delta respecto al puntaje real, R² Test y RMSE."
)

h3("Detalle técnico — Sección expandible")
bullet(
    "Tabla benchmark con todas las métricas: R² Test, R² Train, R² CV (5-fold), CV ± SD, RMSE, MAE."
)
bullet("Gráfico de barras RMSE comparativo (menor = mejor).")
bullet("Gráfico de barras R² Test comparativo (mayor = mejor).")
bullet(
    "Radar multi-métrica normalizado: R² Test, R² CV, R² Train, Precisión RMSE, Precisión MAE."
)
bullet(
    "Nota metodológica: granularidad, n entrenamiento, n test, tipo de split temporal, features utilizadas."
)

h2("5.4 Priorización IPE + IA Gemini (priorizacion.py)")
body(
    "Módulo que calcula el Índice de Priorización Educativa (IPE) para cada institución "
    "en la región seleccionada, permitiendo identificar cuáles requieren intervención "
    "urgente del sistema educativo."
)

h3("Cálculo del IPE")
body(
    "El IPE es un índice compuesto en escala 0–100 (donde 100 = máxima urgencia). "
    "Se construye como suma ponderada de tres componentes, cada uno normalizado "
    "mediante min-max al rango 0–100:"
)
ipe_headers = ["Componente", "Peso", "Variable de origen", "Interpretación"]
ipe_rows = [
    [
        "Deterioro Académico Histórico",
        "40 %",
        "Pendiente regresión lineal de PUNT_GLOBAL por año",
        "Pendiente negativa = mayor riesgo",
    ],
    [
        "Brecha Digital Colectiva",
        "30 %",
        "% estudiantes sin internet en el hogar (FAMI_TIENEINTERNET='No')",
        "Mayor % sin internet = mayor riesgo",
    ],
    [
        "Vulnerabilidad Familiar",
        "30 %",
        "% padres/madres sin educación básica completa",
        "Mayor % vulnerable = mayor riesgo",
    ],
]
add_table(ipe_headers, ipe_rows, [4, 1.5, 5.5, 4.5])

h3("Cuadrantes de acción")
cuad_headers = ["Rango IPE", "Clasificación", "Acción Recomendada"]
cuad_rows = [
    [
        "80–100 pts",
        "🔴 Prioridad Crítica",
        "Planes de choque inmediatos, auditorías pedagógicas, presupuesto de conectividad prioritario",
    ],
    [
        "50–79 pts",
        "🟡 Vulnerabilidad Estructural",
        "Seguimiento semestral, refuerzo tecnológico, programas de padres",
    ],
    [
        "0–49 pts",
        "🟢 Monitoreo Preventivo",
        "Mantenimiento de indicadores, buenas prácticas replicables",
    ],
]
add_table(cuad_headers, cuad_rows, [2.5, 4, 9])

h3("IA Gemini — Diagnóstico por institución")
body(
    "Para cualquier institución del ranking, el usuario puede solicitar un diagnóstico "
    "automático mediante Gemini 2.0 Flash. El prompt estructurado indica al modelo que "
    "actúe como consultor de políticas públicas y genere: (1) interpretación del IPE "
    "de la institución y (2) dos acciones concretas de intervención. El resultado "
    "se muestra en pantalla y puede incluirse en el PDF ejecutivo."
)

h3("PDF Ejecutivo descargable")
body(
    "El módulo genera un reporte PDF usando ReportLab con: encabezado institucional, "
    "tabla del Top 5 de colegios críticos y el diagnóstico IA. El PDF está listo "
    "para presentar ante Gobernaciones o Consejos de Gobierno."
)

h2("5.5 Análisis de Sensibilidad de Pesos del IPE")

h3("¿Qué es el análisis de sensibilidad?")
body(
    "Un análisis de sensibilidad de pesos es una técnica de validación que consiste en "
    "verificar si los resultados de un índice compuesto cambian significativamente cuando "
    "se modifican los pesos asignados a cada componente. En el contexto del IPE, la pregunta "
    "central es: si se cambiaran los pesos 40/30/30 por otras combinaciones razonables, "
    "¿el ranking de instituciones que requieren intervención urgente permanecería igual?"
)
body(
    "Para responder esta pregunta de forma cuantitativa, el sistema implementa la "
    "Correlación de Spearman entre el ranking base (pesos 40/30/30) y el ranking "
    "alternativo generado por cada escenario de pesos. La correlación de Spearman (r) "
    "es una medida no paramétrica que compara dos ordenamientos (rankings) sin asumir "
    "distribución normal: r = 1.0 indica que los dos rankings son idénticos, mientras que "
    "r = 0.0 indica que no hay relación entre ellos."
)

h3("Escenarios de pesos evaluados")
body(
    "Se evaluaron ocho escenarios alternativos, variando cada componente en ±10 puntos "
    "porcentuales mientras los restantes absorben el cambio de forma simétrica:"
)
escenarios_headers = ["Escenario", "Deterioro", "Brecha Digital", "Vulnerab. Familiar", "Descripción"]
escenarios_rows = [
    ["Base", "40 %", "30 %", "30 %", "Configuración de diseño del sistema"],
    ["Det+10", "50 %", "25 %", "25 %", "Mayor énfasis en caída académica"],
    ["Det-10", "30 %", "35 %", "35 %", "Menor énfasis en caída académica"],
    ["Dig+10", "35 %", "40 %", "25 %", "Mayor énfasis en brecha digital"],
    ["Dig-10", "45 %", "20 %", "35 %", "Menor énfasis en brecha digital"],
    ["Vul+10", "35 %", "25 %", "40 %", "Mayor énfasis en vulnerabilidad familiar"],
    ["Vul-10", "45 %", "35 %", "20 %", "Menor énfasis en vulnerabilidad familiar"],
    ["Iguales", "33 %", "33 %", "33 %", "Pesos iguales sin priorización"],
]
add_table(escenarios_headers, escenarios_rows, [2.5, 2, 2.5, 2.5, 6])

h3("Umbral de estabilidad e interpretación")
body(
    "Los resultados se clasifican en tres niveles de estabilidad según el valor de r Spearman:"
)
bullet(
    "Alta estabilidad (r >= 0.95): el ranking cambia en menos del 5% de los puestos relativos. "
    "Los pesos son intercambiables dentro del rango evaluado y la eleccion de diseno es robusta."
)
bullet(
    "Estabilidad moderada (0.85 <= r < 0.95): existe variacion apreciable pero las instituciones "
    "criticas siguen apareciendo en los primeros puestos. Se recomienda documentar la justificacion "
    "de la eleccion de pesos con referencias bibliograficas."
)
bullet(
    "Baja estabilidad (r < 0.85): el ranking cambia significativamente. Indica que los pesos "
    "exactos son determinantes del resultado y deben derivarse empiricamente (regresion OLS o AHP)."
)

h3("Justificacion de los pesos 40/30/30")
body(
    "La eleccion de los pesos 40 % deterioro / 30 % brecha digital / 30 % vulnerabilidad familiar "
    "sigue una logica de decision multicriterio basada en tres principios:"
)
bullet(
    "Primacia del resultado academico (40 %): el deterioro del puntaje PUNT_GLOBAL es el indicador "
    "mas directo de la necesidad de intervencion pedagogica. Una institucion con caida sistematica "
    "en resultados requiere atencion inmediata independientemente de sus condiciones contextuales. "
    "Por ello recibe el mayor peso individual."
)
bullet(
    "Equidad entre factores de contexto (30 % + 30 %): la brecha digital y la vulnerabilidad familiar "
    "son factores estructurales que explican causalmente el deterioro academico, pero ninguno "
    "es determinante por si solo. Asignarles el mismo peso refleja que ambas dimensiones tienen "
    "igual relevancia desde la perspectiva de politica publica: la conectividad afecta el acceso "
    "a recursos de aprendizaje, mientras que la educacion de los padres afecta el capital cultural "
    "y el acompanamiento en el hogar."
)
bullet(
    "Suma 100 % con integrantes >= 25 %: se garantiza que ningun componente quede marginado. "
    "Un peso inferior al 25 % en cualquier dimension haria que esa dimension fuera practicamente "
    "irrelevante en el calculo, lo que no refleja la evidencia empirica sobre los determinantes "
    "del rendimiento en Colombia (ver estudios ICFES sobre NSE y resultados Saber 11)."
)
body(
    "El analisis de sensibilidad implementado en el dashboard (seccion expandible del modulo "
    "Priorizacion IPE) permite verificar empiricamente que este diseno 40/30/30 es robusto: "
    "si la correlacion de Spearman entre el ranking base y los rankings alternativos supera 0.95 "
    "en todos los escenarios, la eleccion de pesos queda validada como una decision de diseno "
    "justificable y no arbitraria."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. COVID-19
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Vista COVID-19 — Análisis del Impacto Pandémico")
body(
    "La pestaña COVID-19 (covid.py) es una vista dedicada al análisis del impacto "
    "de la pandemia del coronavirus en el rendimiento académico de los estudiantes "
    "colombianos de grado 11. El análisis divide la serie temporal en tres períodos "
    "bien definidos:"
)
periodos_headers = ["Período", "Años", "Modalidad educativa"]
periodos_rows = [
    ["Pre-COVID", "2015–2019", "Presencial — condiciones normales"],
    ["Pandemia (COVID)", "2020–2021", "Virtual / Alternancia / Cierre de colegios"],
    ["Post-COVID", "2022–2025+", "Retorno gradual a presencialidad"],
]
add_table(periodos_headers, periodos_rows, [3, 3, 9.5])

h2("6.1 KPIs de Impacto")
body(
    "Se calculan cinco métricas de resumen desde los datos reales o el mock paramétrico:"
)
bullet("Promedio global Pre-COVID: media del puntaje global en el período 2015–2019.")
bullet("Promedio global COVID: media del puntaje global en 2020–2021.")
bullet("Promedio global Post-COVID: media del puntaje global en 2022 en adelante.")
bullet("Caída pandemia: diferencia en puntos entre el promedio Pre-COVID y COVID.")
bullet(
    "Índice de Recuperación: porcentaje de la caída que ha sido recuperada en el período Post-COVID, calculado como (Post−COVID) / (Pre−COVID) × 100, truncado al 100%."
)

h2("6.2 Evolución Temporal con Banda COVID")
body(
    "Gráfico de líneas coloreadas por período (azul Pre-COVID, rojo COVID, verde "
    "Post-COVID) con una banda rectangular sombreada sobre los años 2020–2021 que "
    "señala visualmente el período de pandemia. Los marcadores de cada año permiten "
    "ver el valor exacto en hover."
)

h2("6.3 Análisis por Área de Conocimiento")
body(
    "Gráfico de barras agrupadas donde el eje X son las áreas (Global, Matemáticas, "
    "Lectura Crítica, Ciencias, Sociales, Inglés) y cada grupo tiene tres barras "
    "correspondientes a los tres períodos. Permite identificar cuáles áreas fueron "
    "más afectadas por la pandemia y cuáles muestran mejor recuperación."
)

h2("6.4 Caída y Recuperación por Departamento")
body("Dos gráficos horizontales en paralelo para los 33 departamentos de Colombia:")
bullet(
    "Caída en puntos (Pre → COVID): escala de color verde a rojo indicando la magnitud de la pérdida."
)
bullet(
    "Índice de Recuperación Post-COVID (%): escala de color rojo a verde indicando qué tan cerca está cada departamento de volver al nivel pre-pandemia."
)

h2("6.5 Brecha de Género por Período")
body(
    "Gráfico de barras agrupadas comparando el promedio global de estudiantes "
    "masculinos y femeninos en cada período. Permite detectar si la pandemia amplificó "
    "o redujo la brecha de género preexistente."
)

h2("6.6 Tabla de Insights Accionables")
body(
    "Tabla con seis indicadores críticos del impacto pandémico, incluyendo la "
    "situación observada y la acción recomendada para cada uno, con codificación "
    "de prioridad (Crítica / Alta / Media)."
)

h2("6.7 Calculadora Interactiva de Recuperación")
body(
    "Sección expandible donde el coordinador puede ingresar los puntajes pre-COVID, "
    "durante COVID y actual de su colegio, junto con una meta de recuperación (%); "
    "el sistema calcula automáticamente el índice de recuperación actual y los puntos "
    "faltantes para alcanzar la meta."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. MÓDULO IA — GEMINI
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Módulo de Inteligencia Artificial — Gemini")
body(
    "El módulo de IA (src/icfes/dashboard/ai/) centraliza toda la interacción con "
    "la API de Gemini 2.0 Flash de Google. Está compuesto por un cliente centralizado "
    "y un sistema de prompts estructurados por módulo y destinatario."
)

h2("7.1 Cliente Gemini (ai/client.py)")
body(
    "El cliente gestiona la autenticación con la API key buscándola en dos fuentes "
    "en orden de prioridad: (1) st.secrets['GEMINI_API_KEY'] del archivo "
    ".streamlit/secrets.toml y (2) variable de entorno GEMINI_API_KEY. "
    "Si ninguna está configurada, los botones de IA quedan deshabilitados sin romper "
    "el resto del dashboard. La función generate(prompt) ejecuta la llamada a la API "
    "y retorna el texto generado."
)

h2("7.2 Prompts Estructurados")
body(
    "Los prompts se construyen programáticamente con la función build_*() "
    "correspondiente a cada caso de uso. Cada prompt tiene cinco secciones:"
)
bullet(
    "ROL: define la identidad del modelo (ej: 'Eres un consultor de política pública educativa')."
)
bullet("CONTEXTO APP: describe brevemente el sistema EduMetrics Saber11.")
bullet(
    "DESTINATARIO: indica quién recibirá la respuesta (Rector, Secretario, Estudiante)."
)
bullet(
    "DATOS ANALIZADOS: inyecta los valores numéricos específicos de la institución/estudiante."
)
bullet("TIPO RESPUESTA: instruye el formato esperado (extensión, tono, estructura).")

h2("7.3 Casos de Uso de IA")
ia_headers = ["Módulo", "Función build_*()", "Rol del modelo", "Destinatario", "Output"]
ia_rows = [
    [
        "Priorización IPE",
        "build_intervencion_ipe()",
        "Consultor políticas públicas",
        "Secretario de Educación",
        "Diagnóstico IPE + 2 acciones de intervención",
    ],
    [
        "Perfilamiento colectivo",
        "build_colectivo()",
        "Consultor política educativa",
        "Rector / Secretaría",
        "Justificación del perfil institucional + 2 recomendaciones pedagógicas",
    ],
    [
        "Perfilamiento individual",
        "build_individual()",
        "Orientador vocacional",
        "Estudiante grado 11",
        "¿Por qué este perfil? + carrera a explorar",
    ],
]
add_table(ia_headers, ia_rows, [3.5, 3.5, 3.5, 3.5, 5])
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. MODELOS ML — SIMULADOR
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Modelos de Machine Learning — Simulador Predictivo")

h2("8.1 Proceso de Entrenamiento (scripts/train_simulador.py)")
body(
    "El entrenamiento se ejecuta independientemente del dashboard mediante el comando "
    "uv run python scripts/train_simulador.py. El proceso sigue estos pasos:"
)
bullet(
    "Extracción de datos: consulta DuckDB sobre los Parquet para obtener promedios agregados a granularidad colegio × municipio × año."
)
bullet(
    "Variables predictoras (features): pct_internet (% con internet), pct_educacion_sup (% padres con educación superior), promedio_estrato."
)
bullet(
    "Variable objetivo: punt_global (promedio del puntaje global por colegio/municipio/año)."
)
bullet(
    "Split temporal honesto: test set = últimos 2 años disponibles. Esto evita data leakage ya que los datos futuros no se usan en entrenamiento."
)
bullet("Validación cruzada: 5-fold sobre el train set para estimar la generalización.")
bullet(
    "Reentrenamiento final: los tres modelos se reentrenan con TODOS los datos (train + test) antes de guardarse en el pkl."
)

h2("8.2 Los Tres Modelos")
models_headers = ["Modelo", "Clave", "Hiperparámetros clave", "Característica"]
models_rows = [
    [
        "Ridge Regression",
        "ridge",
        "alpha=50.0, StandardScaler",
        "Modelo lineal con regularización L2. Rápido, interpretable, robusto ante multicolinealidad. Coeficientes directamente interpretables.",
    ],
    [
        "Random Forest Regressor",
        "rf",
        "n_estimators=200, max_depth=12, min_samples_leaf=5, random_state=42",
        "Ensemble de árboles de decisión. Captura relaciones no lineales. Robusto ante outliers. Alta capacidad de ajuste.",
    ],
    [
        "Gradient Boosting Regressor",
        "gbr",
        "n_estimators=300, learning_rate=0.05, max_depth=4, subsample=0.8, random_state=42",
        "Boosting secuencial. Minimiza el error residual iterativamente. Mayor complejidad, mejor performance en muchos casos.",
    ],
]
add_table(models_headers, models_rows, [3.5, 2, 4.5, 7])

h2("8.3 Métricas de Evaluación")
body(
    "Para cada modelo se calculan y almacenan en el bundle .pkl las siguientes métricas:"
)
metrics_headers = ["Métrica", "Símbolo", "Interpretación"]
metrics_rows = [
    [
        "Coeficiente de determinación (test)",
        "R² Test",
        "Proporción de varianza explicada en el conjunto de prueba. Mayor = mejor.",
    ],
    [
        "Coeficiente de determinación (train)",
        "R² Train",
        "Ajuste sobre los datos de entrenamiento. Muy superior a R² Test indica sobreajuste.",
    ],
    [
        "R² Validación Cruzada (5-fold)",
        "R² CV",
        "Estimación robusta de la capacidad de generalización.",
    ],
    [
        "Desviación estándar del CV",
        "CV ± SD",
        "Estabilidad del modelo entre los 5 folds. Menor = más estable.",
    ],
    [
        "Error cuadrático medio",
        "MSE",
        "Promedio de los errores al cuadrado. Penaliza errores grandes.",
    ],
    [
        "Raíz del error cuadrático medio",
        "RMSE",
        "Interpretable en las mismas unidades del puntaje (puntos ICFES).",
    ],
    [
        "Error absoluto medio",
        "MAE",
        "Promedio del valor absoluto de los errores. Más robusto ante outliers.",
    ],
]
add_table(metrics_headers, metrics_rows, [4.5, 2.5, 9])

h2("8.4 Selección del Mejor Modelo")
body(
    "El mejor modelo es aquel con el mayor R² Test. Se marca con el badge '🏆 Mejor modelo' "
    "en el dashboard y sus resultados aparecen destacados con borde de color y efecto "
    "glow visual. En el expander de detalle técnico, el gráfico de barras destaca el "
    "modelo ganador con opacidad completa mientras los demás aparecen en 45% de opacidad."
)

h2("8.5 Bundle .pkl — Estructura")
body(
    "El archivo models/simulador_models.pkl contiene el bundle completo con los tres "
    "modelos y toda la metadata necesaria para reproducir los resultados:"
)
bullet("models: dict con los tres pipelines entrenados (ridge, rf, gbr).")
bullet(
    "models_meta: dict con label, short, color, icon y todas las métricas por modelo."
)
bullet("best_model_key: clave del modelo con mayor R² Test.")
bullet("granularity: descriptor del nivel de agregación ('colegio/municipio/año').")
bullet("n_train, n_test, n_total: tamaños de los conjuntos de datos.")
bullet("features: lista de variables predictoras.")
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. ESTRUCTURA TÉCNICA
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Estructura Técnica del Proyecto")

h2("9.1 Árbol de Directorios")
body("La siguiente tabla describe cada directorio y archivo relevante del proyecto:")
struct_headers = ["Ruta", "Descripción"]
struct_rows = [
    ["files/raw/", "Archivos planos ICFES originales (.txt con separador ;)"],
    ["files/parquet/", "Archivos .parquet comprimidos (salida del ETL, ~snappy)"],
    ["files/dane/", "Catálogo DIVIPOLA del DANE (códigos → nombres depto/municipio)"],
    ["models/", "Bundle de modelos ML entrenados (simulador_models.pkl)"],
    ["queries/", "Archivos SQL reutilizables con placeholder {parquet}"],
    ["scripts/train_simulador.py", "Script de entrenamiento de los 3 modelos ML"],
    [
        "src/icfes/settings.py",
        "Configuración central: variables de entorno, rutas, backend",
    ],
    ["src/icfes/logger.py", "Logger configurado con formato estándar"],
    ["src/icfes/etl/config.py", "Mapeos de columnas históricos (MAP_ATTR_ICFES)"],
    ["src/icfes/etl/pipeline.py", "Ejecutor del flujo ETL completo"],
    ["src/icfes/etl/schemas.py", "Esquemas estrictos PyArrow"],
    ["src/icfes/core/query_service.py", "Abstracción DuckDB con soporte multi-backend"],
    ["src/icfes/api/main.py", "FastAPI stub (Fase 2, pendiente)"],
    [
        "src/icfes/dashboard/app.py",
        "Orquestador: configuración, tabs, routing, sidebar",
    ],
    [
        "src/icfes/dashboard/components/theme.py",
        "CSS global oscuro y función dark_layout()",
    ],
    [
        "src/icfes/dashboard/components/sidebar.py",
        "Filtros globales → (where_clause, filtros...)",
    ],
    ["src/icfes/dashboard/components/animations.py", "Partículas animadas + anime.js"],
    [
        "src/icfes/dashboard/ai/client.py",
        "Cliente Gemini centralizado + gestión de API key",
    ],
    ["src/icfes/dashboard/ai/prompts/", "Plantillas de prompts por módulo"],
    [
        "src/icfes/dashboard/pages/inicio.py",
        "Pantalla de bienvenida con cards de navegación",
    ],
    [
        "src/icfes/dashboard/pages/analisis.py",
        "KPIs + tendencia + radar + mapa departamentos",
    ],
    [
        "src/icfes/dashboard/pages/tendencias.py",
        "Tendencias históricas + proyección OLS",
    ],
    [
        "src/icfes/dashboard/pages/coordinador.py",
        "Brechas institucionales + impacto socioeconómico",
    ],
    [
        "src/icfes/dashboard/pages/secretario.py",
        "Equidad regional Oficial/Privado + Urbano/Rural",
    ],
    ["src/icfes/dashboard/pages/simulador.py", "Simulador ML multi-modelo + benchmark"],
    [
        "src/icfes/dashboard/pages/priorizacion.py",
        "IPE + ranking instituciones + IA Gemini + PDF",
    ],
    [
        "src/icfes/dashboard/pages/covid.py",
        "Análisis impacto pandemia Pre/Durante/Post",
    ],
    [
        "src/icfes/dashboard/pages/perfilamiento.py",
        "Perfiles vocacionales + IA Gemini (oculto en menú)",
    ],
    ["pyproject.toml", "Declaración de dependencias y metadatos del proyecto (uv)"],
    ["Makefile", "Comandos abreviados: etl, dashboard, api, lint, test, clean"],
    [".env / .env.example", "Variables de entorno: backend, rutas, credenciales cloud"],
    [".streamlit/secrets.toml", "API key de Gemini para el dashboard"],
]
add_table(struct_headers, struct_rows, [6.5, 9])

h2("9.2 Backends de Almacenamiento")
body(
    "El QueryService soporta tres backends intercambiables via la variable de entorno "
    "STORAGE_BACKEND, sin cambios de código en el dashboard:"
)
backend_headers = ["Backend", "Valor", "Descripción", "Configuración requerida"]
backend_rows = [
    [
        "Local (default)",
        "local",
        "Parquet en files/parquet/ del filesystem local",
        "Solo PARQUET_PATH",
    ],
    [
        "AWS S3",
        "s3",
        "Parquet en bucket S3 vía DuckDB httpfs",
        "AWS_REGION, AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_PARQUET_PATH",
    ],
    [
        "Supabase",
        "supabase",
        "Supabase Storage S3-compatible con path-style URLs",
        "SUPABASE_S3_ENDPOINT, SUPABASE_ACCESS_KEY, SUPABASE_SECRET_KEY, SUPABASE_PARQUET_PATH",
    ],
]
add_table(backend_headers, backend_rows, [2.5, 2, 4.5, 7])
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. ARQUITECTURA Y DINÁMICA
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Arquitectura y Dinámica de Módulos")

h2("9.1 Flujo de Datos End-to-End")
body("El flujo completo de datos en la plataforma sigue el siguiente recorrido:")
bullet(
    "1. Fuente: Archivos planos .txt del ICFES (separador ';', ~4.3 GB total, 22 semestres 2015–2025)."
)
bullet(
    "2. ETL: pipeline.py lee, mapea columnas, limpia texto, enriquece con DIVIPOLA e infiere tipos."
)
bullet(
    "3. Almacenamiento: Parquet con compresión Snappy en files/parquet/ (un archivo por semestre)."
)
bullet(
    "4. QueryService: abstracción DuckDB que resuelve el placeholder {parquet} → read_parquet('files/parquet/*.parquet') en cada SQL."
)
bullet(
    "5. Caché de recursos: @st.cache_resource para el QueryService (singleton) y @st.cache_data(ttl=600) para resultados de queries frecuentes."
)
bullet(
    "6. Dashboard: Streamlit renderiza KPIs, gráficos Plotly y componentes interactivos en respuesta a interacciones del usuario."
)
bullet(
    "7. IA: cuando se solicita, el texto extraído de DuckDB se inyecta en el prompt de Gemini y el resultado se muestra en pantalla."
)

h2("9.2 Dinámica del Sidebar y Filtros Globales")
body(
    "El componente sidebar.py gestiona los filtros globales que afectan la vista Explorar. "
    "Devuelve una tupla (where_clause, sel_anos, sel_deptos, sel_genero, sel_naturaleza). "
    "La where_clause es una cadena SQL que comienza con 'WHERE' cuando hay filtros activos, "
    "y se pasa directamente a las queries de analisis.py mediante interpolación de f-strings. "
    "El sidebar se abre automáticamente cuando el usuario navega a la sección Análisis "
    "(controlado por JavaScript inyectado en la página), y se cierra en las demás vistas."
)

h2("9.3 Sistema de Caché Multi-nivel")
body(
    "Para garantizar un buen rendimiento sobre un volumen grande de datos, el sistema "
    "implementa caché en dos niveles:"
)
bullet(
    "@st.cache_resource — QueryService: singleton que mantiene la conexión DuckDB abierta entre rerenders de Streamlit."
)
bullet(
    "@st.cache_data(ttl=600) — Resultados de queries: caché de 10 minutos para listas de municipios, colegios, indicadores y cálculos IPE, evitando re-ejecutar queries costosas."
)

h2("9.4 Manejo de Datos Faltantes y Casos Especiales")
body("El sistema implementa varias estrategias para garantizar robustez:")
bullet(
    "NaN en punt_ingles: columna con alta tasa de datos faltantes. En tendencias → tratado como 0 (no excluye la fila). En perfilamiento → tratado como NULL (AVG lo ignora)."
)
bullet(
    "ano como float: DuckDB puede retornar años como 2015.0. Todas las queries aplican CAST(ano AS INTEGER)."
)
bullet(
    "Datos mock: cuando una query falla o retorna vacío, se activa un fallback con datos paramétricos realistas para que el dashboard sea funcional sin datos reales."
)
bullet(
    "Try/except generalizado: cada sección gráfica captura excepciones individualmente para que un error en un componente no bloquee el resto de la página."
)

h2("9.5 Gestión de Dependencias con uv")
body(
    "El proyecto usa uv como gestor de paquetes moderno (Astral). El archivo "
    "pyproject.toml declara las dependencias con grupos opcionales:"
)
bullet("Grupo base: duckdb, pyarrow, python-dotenv, loguru.")
bullet(
    "Grupo [dashboard]: streamlit, plotly, streamlit-option-menu, scipy, scikit-learn, reportlab, joblib."
)
bullet("Grupo [etl]: pyarrow, pandas (si se instala separadamente del base).")
body(
    "El entorno se activa automáticamente con uv run, que maneja el virtualenv "
    "(.venv/) de forma transparente sin necesidad de activación manual."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Conclusiones")
body(
    "EduMetrics Saber11 demuestra que es posible construir una plataforma de analítica "
    "educativa completa, de alto rendimiento y con IA generativa integrada, utilizando "
    "exclusivamente herramientas de código abierto y una arquitectura bien estructurada. "
    "Los principales logros del proyecto son:"
)
bullet(
    "ETL robusto: procesamiento automático de 22 semestres de microdatos con normalización de esquemas históricos, codificación, DIVIPOLA y compresión, reduciendo el tamaño de datos en ~75%."
)
bullet(
    "Análisis EDA completo: exploración interactiva con KPIs, radar de competencias, evolución temporal, distribución de puntajes y proyección OLS post-COVID, todo con filtros globales en tiempo real."
)
bullet(
    "Herramientas diferenciadas por usuario: los tableros de Coordinador, Secretario, Simulador y Priorización están diseñados específicamente para el flujo de trabajo de cada actor del sistema educativo."
)
bullet(
    "Machine Learning riguroso: evaluación honesta con split temporal, CV 5-fold y comparación simultánea de tres modelos, evitando sobreajuste y data leakage."
)
bullet(
    "IPE — Índice innovador: el Índice de Priorización Educativa ofrece una métrica compuesta y objetiva para priorizar la intervención pública con base en datos."
)
bullet(
    "IA Generativa integrada: el módulo Gemini convierte datos numéricos en diagnósticos textuales accionables, democratizando el análisis para usuarios sin formación estadística."
)
bullet(
    "Análisis COVID-19: cuantificación del impacto pandémico por área, departamento y género, con calculadora interactiva de recuperación para coordinadores."
)
bullet(
    "Escalabilidad: la arquitectura desacoplada (ETL → Parquet → QueryService → Dashboard) permite migrar a cloud (S3/Supabase) sin modificar el código de la capa de presentación."
)

doc.add_paragraph()
divider()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run(
    f"EduMetrics Saber11 — Informe Técnico · Generado el {datetime.date.today().strftime('%d/%m/%Y')} · Python + DuckDB + Streamlit"
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x71, 0x88, 0xA8)
r.italic = True


# ── Guardar ───────────────────────────────────────────────────────────────────
OUTPUT = "Informe_EduMetrics_Saber11.docx"
doc.save(OUTPUT)
print(f"[OK] Informe generado: {OUTPUT}")
